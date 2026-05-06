#!/usr/bin/env python3
"""
generate_brief_v3.py - 参考铝产业简报样式的行业简报生成器

完全依照"一周全球铝产业资讯分析简报.docx"的样式体系：
  字体：微软雅黑
  封面：深橙色(B45309)/深灰(1F2937)大标题 + 灰色(6B7280)信息行
  章节标题：B45309 16pt 加粗  |  小节标题：1F2937 12pt 加粗
  表格表头：#1E3A5F 白字 10pt 加粗
  [利好]#D1FAE5 [利空]#FEE2E2 [关注]#EDE9FE
  [紧急]#FEE2E2 [重要]#FEF3C7 [常规]#DBEAFE [长期]#EDE9FE

用法：
  python3 generate_brief_v3.py --date 2026-04-29
  python3 generate_brief_v3.py --date 2026-04-29 --dept 工业互联网处
  python3 generate_brief_v3.py --date 2026-04-29 --skip-llm
"""

import json, os, sys, re, time, random, requests
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Optional
from urllib.parse import urlparse

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, "data")
CONFIG_DIR = os.path.join(BASE_DIR, "config")

LLM_API_URL = "https://api.deepseek.com/v1/chat/completions"
LLM_API_KEY = "sk-xxx"
LLM_MODEL = "deepseek-v4-flash"

# ===================== 配色系统（参考铝产业简报） =====================
COLOR_PRIMARY     = RGBColor(0x1E, 0x3A, 0x5F)  # 表格表头底色
COLOR_ACCENT      = RGBColor(0xB4, 0x53, 0x09)  # 章节标题/标签(橙色)
COLOR_DARK        = RGBColor(0x1F, 0x29, 0x37)  # 正文深灰色
COLOR_GRAY        = RGBColor(0x6B, 0x72, 0x80)  # 副标题/信息灰色
COLOR_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LINK        = RGBColor(0x25, 0x63, 0xEB)  # 链接蓝色

# 舆情标签配色
COLOR_LIHOU   = "D1FAE5"  # 利好
COLOR_LIKONG  = "FEE2E2"  # 利空
COLOR_GUANZHU = "EDE9FE"  # 关注

# 优先级配色
PRIO_COLORS = {
    "紧急": "FEE2E2",
    "重要": "FEF3C7",
    "常规": "DBEAFE",
    "长期": "EDE9FE",
}

FONT_NAME = "Microsoft YaHei"

# ===================== 资讯类型关键词分类 =====================
INFO_TYPE_KEYWORDS = {
    "政策": ["政策", "规划", "监管", "印发", "通知", "公告", "指导意见", "办法", "方案",
             "条例", "法规", "法律", "标准", "规定", "发文", "工信部", "发改委",
             "国务院", "国家", "省", "市", "厅", "局", "委"],
    "市场": ["市场", "价格", "行情", "报价", "走势", "交易", "供需", "库存",
             "指数", "采购", "销售", "贸易", "出口", "进口", "需求", "供应",
             "涨幅", "跌幅", "波动", "涨价", "降价", "成交"],
    "企业": ["企业", "公司", "集团", "股份", "有限", "华为", "阿里", "腾讯",
             "投资", "产能", "扩建", "投产", "并购", "融资", "上市"],
    "技术": ["技术", "创新", "突破", "研发", "专利", "芯片", "AI", "算法",
             "智能", "数字化", "自动化", "新材料", "工艺", "设备"],
    "舆情": ["问题", "整改", "处罚", "违规", "投诉", "事故", "安全", "环保",
             "负面", "曝光", "调查", "风险", "警示", "罚款", "停产"],
}

SENTIMENT_LABEL_KEYWORDS = {
    "利好": ["增长", "利好", "突破", "提升", "扩大", "投产", "创新", "获批",
             "支持", "补贴", "表彰", "领先", "盈利", "上升"],
    "利空": ["下跌", "利空", "处罚", "违规", "下降", "亏损", "停产", "事故",
             "诉讼", "罚款", "查封", "限产", "减产", "关闭"],
    "关注": ["关注", "注意", "风险", "警惕", "波动", "调整", "不确定", "变化",
             "争议", "存疑", "审查", "关注"],
}

# ===================== 评分系统 =====================
def score_article(item: dict, dept_name: str) -> float:
    title = item.get('title', '')
    abstract = item.get('abstract', item.get('summary', ''))
    source = item.get('source_name', item.get('source', ''))
    date_str = item.get('date', '')
    link = item.get('link', '')
    combined = (title + ' ' + abstract).lower()
    score = 0.0

    if date_str:
        try:
            d = datetime.strptime(date_str[:10], "%Y-%m-%d")
            days_diff = (datetime.now() - d).days
            if days_diff <= 1: score += 5.0
            elif days_diff <= 3: score += 4.0
            elif days_diff <= 7: score += 3.0
            elif days_diff <= 14: score += 2.0
            elif days_diff <= 30: score += 1.0
        except: score += 1.0
    else: score += 0.5

    source_lower = source.lower()
    url_lower = link.lower()
    gov_markers = ['gov.cn', 'miit.gov', 'stats.gov', 'ndrc', 'nhc.gov',
                   'samr.gov', 'mnr.gov', 'mee.gov', 'people.com', 'cctv',
                   '工信部', '国家', '国务院', '发改委', '统计局']
    assoc_markers = ['协会', '学会', '研究院', '研究所', '联盟', '委员会',
                     'caict', 'ccid', 'csia', 'cie', 'cnea', 'cvida',
                     'chinania', 'aladdiny', 'baiinfo', 'smm', 'yyjjb',
                     'cntcm', '100ppi', 'cbea', 'oilchem']
    for m in gov_markers:
        if m.lower() in source_lower or m.lower() in url_lower:
            score += 5.0
            break
    else:
        for m in assoc_markers:
            if m.lower() in source_lower or m.lower() in url_lower:
                score += 3.0
                break
        else: score += 1.0

    dept_kw_map = {
        "工业互联网处": ["电子信息", "半导体", "芯片", "集成电路", "工业互联网",
                        "5G", "6G", "AI", "大数据", "云计算", "物联网", "数字化",
                        "智能制造", "电子制造", "电子元件", "传感器", "数据中心",
                        "算力", "光伏", "锂电", "电池", "储能"],
        "医药处": ["医药", "药品", "医疗", "医保", "集采", "中医药", "中药",
                   "创新药", "生物", "疫苗", "医疗器械", "临床", "药企", "制药"],
        "原材料_金属_非金属处": ["铝", "磷", "矿产", "有色金属", "钢铁", "氧化铝",
                                 "电解铝", "碳酸锂", "电池材料", "稀土", "钢材",
                                 "产能", "大宗商品", "新材料"],
    }
    kw_list = dept_kw_map.get(dept_name, [])
    matched = sum(1 for kw in kw_list if kw.lower() in combined)
    if matched >= 3: score += 5.0
    elif matched >= 2: score += 4.0
    elif matched >= 1: score += 2.0

    if len(title) >= 20: score += 3.0
    elif len(title) >= 12: score += 2.0
    else: score += 1.0
    if len(title) < 15 and matched == 0: score -= 2.0

    ad_kw = ['广告', '推广', '促销', '招聘', '转让', '出租']
    for kw in ad_kw:
        if kw in combined:
            score -= 10.0
            break
    return score


def deduplicate_articles(articles: List[dict]) -> List[dict]:
    seen = []
    result = []
    for a in articles:
        key = a.get('title', '')[:15]
        is_dup = any(key in s or s in key for s in seen)
        if not is_dup:
            seen.append(key)
            result.append(a)
    return result


# ===================== 原文获取 =====================
def _is_garbage_text(text: str) -> bool:
    """检测反爬页面、无意义页脚、版权声明等垃圾内容"""
    garbage_signals = [
        "安全检查", "请稍候", "安全验证", "正在验证", "loading", "challenge",
        "版权所有", "Copyright", "All Rights Reserved", "沪ICP备", "浙B2-",
        "粤ICP备", "京ICP备", "ICP证", "免责声明", "友情链接",
        "关于我们", "联系我们", "网站地图", "网站声明", "帮助中心",
        "广告服务", "合作伙伴", "招贤纳士", "诚聘英才",
        "页面加载中", "数据加载中", "加载失败",
    ]
    text_lower = text.lower()
    # 统计垃圾信号命中次数
    hit_count = sum(1 for g in garbage_signals if g.lower() in text_lower)
    # 如果命中 >= 2 个垃圾信号，或命中"安全检查"等强信号，视为垃圾
    if hit_count >= 2:
        return True
    if any(s in text_lower for s in ["安全检查", "请稍候", "安全验证"]):
        return True
    # 检查是否主要是版权/ICP信息（短文本且只包含这些）
    if len(text) < 100 and ("版权所有" in text_lower or "icp" in text_lower):
        return True
    return False


def _clean_summary(text: str) -> str:
    """清理摘要中的噪声文字"""
    if not text:
        return ""
    # 1. LLM摘要标记
    text = re.sub(r'\*\*(?:专业|核心)?摘要\s*[：:]?\s*\*\*\s*', '', text)
    text = re.sub(r'^摘要[：:]\s*', '', text)

    # 2. 导航面包屑："首页 > 新闻发布 > 部领导活动" 及类似分隔符
    text = re.sub(r'^(?:\s*[>＞]\s*[\u4e00-\u9fff\w]+\s*)+', '', text)
    text = re.sub(r'(?:\s*[>＞]\s*[\u4e00-\u9fff\w]+\s*)+$', '', text)
    text = re.sub(r'(?:首页|当前位置|您现在的位置|您的位置)[\s　]*[：:>＞]\s*', '', text)
    text = re.sub(r'(?<=\S)\s+(?:[\u4e00-\u9fff\w]{2,10}\s*[>＞]\s*){2,}(?:[\u4e00-\u9fff\w]{2,10})\s+', ' ', text)
    text = re.sub(r'^(?:[\u4e00-\u9fff]{2,6}\s+(?=[\u4e00-\u9fff]{6,}))', '', text)
    text = re.sub(r'\s*[>＞]\s*[\u4e00-\u9fff\w]*\s*', ' ', text)

    # 3. 发布时间
    text = re.sub(r'发布时间[：:]\s*\d{4}[-/年]\d{1,2}[-/月]\d{1,2}\s*\d{0,2}[:：]?\d{0,2}\s*', '', text)
    text = re.sub(r'发布时间：\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}\s+来源：[^\s]+\s*', '', text)
    text = re.sub(r'来源[：:][\s\S]*?(?=[\u4e00-\u9fa5]{4,}\s|\u3000|\s{2,}|$)', '', text)
    text = re.sub(r'来源：[^\s]+\s*', '', text)
    text = re.sub(r'来源[：:][^\n]*', '', text)
    text = re.sub(r'(?:中国网|工信微报|新华网|人民网|央视|经济日报|澎湃新闻|新浪财经|网易新闻|腾讯新闻|凤凰网|观察者网|每日经济新闻|第一财经|21世纪经济报道|证券日报|证券时报|上海证券报|中国证券报|经济参考报|科技日报|光明日报|法制日报|中国青年报|北京青年报|南方日报|南方都市报|羊城晚报|广州日报|深圳特区报|新京报|北京日报|解放日报|文汇报|新民晚报|杭州日报|浙江日报|江苏日报|四川日报|贵州日报|云南日报|湖北日报|湖南日报|河南日报|河北日报|山东日报|山西日报|辽宁日报|吉林日报|黑龙江日报|福建日报|江西日报|安徽日报|陕西日报|甘肃日报|青海日报|海南日报|内蒙古日报|新疆日报|西藏日报|宁夏日报|广西日报)[\s　]*\d{4}年\d{1,2}月\d{1,2}日', '', text)

    # 4. 公文头信息
    text = re.sub(r'(?:发文机关|标\s*题|发文字号|成文日期|发布日期|发布机构|分\s*类)[\s　]*[：:][^\n]*\n?', '', text)

    # 5. 去掉正文前/后的页面装饰文字和页脚垃圾
    text = re.sub(r'(?:字号|字体)[\s　]*[：:][\s　]*\[?\s*[\u4e00-\u9fff]*(?:\s*[\u4e00-\u9fff]*)*\s*\]?[\s　]*', '', text)
    text = re.sub(r'视力保护色[：:][\s　]*', '', text)
    text = re.sub(r'\[\s*大\s*中\s*小\s*\]', '', text)
    text = re.sub(r'【返回顶部】\s*【关闭窗口】\s*【打印本页】', '', text)
    text = re.sub(r'【返回顶部】', '', text)
    text = re.sub(r'【关闭窗口】', '', text)
    text = re.sub(r'【打印本页】', '', text)
    text = re.sub(r'相关解读.*?(?:$|\n)', '', text)
    text = re.sub(r'政策解读[：:][^\n]*', '', text)
    text = re.sub(r'一图读懂[：:][^\n]*', '', text)
    text = re.sub(r'\[打印\]|\[关闭\]', '', text)
    text = re.sub(r'分享到[：:][^\n]*', '', text)
    text = re.sub(r'免责声明[^\n]*', '', text)
    text = re.sub(r'版权所有[^\n]*', '', text)
    text = re.sub(r'(?:京|沪|粤|浙|苏|川|黔)\w{0,10}备[^\s]*', '', text)
    text = re.sub(r'(?:工业和信息化部|网站)\s*标识[^\n]*', '', text)

    # 6. 去掉多余的空白
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def fetch_article_text(url: str, timeout: int = 15) -> str:
    if not url or url == '#':
        return ""
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8"}
        resp = requests.get(url, headers=headers, timeout=timeout)
        resp.encoding = 'utf-8'
        # HTTP 请求被拒绝（307/403/412）时自动重试 HTTPS
        if resp.status_code in (307, 403, 412) and url.startswith('http://'):
            https_url = 'https://' + url[7:]
            resp = requests.get(https_url, headers=headers, timeout=timeout)
            resp.encoding = 'utf-8'
        html = resp.text
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "lxml")
        for tag in soup.find_all(['script', 'style', 'nav', 'footer', 'header',
                                   'aside', 'noscript', 'iframe']):
            tag.decompose()
        content_selectors = [
            '#Zoom', '.trs_editor_view', '.TRS_Editor', '.TRS_PreAppend',
            '.Custom_UnionStyle', '.TRS_UEDITOR', '.article-content-inner',
            '#content', '#article', '#context', '.content', '.article',
            '.news-content', '.news-detail', '.detail-content',
            '.con', '.body', '.zw', '.text',
            '#con_con', '.ccontent', '.cmain', '.wzy-wrapper',
            'article', 'main',
            '.post-content', '.entry-content', '.article-content',
            '.text-content', '.page-content', '.main-content', '.detail',
            '.info-content', '.left',
        ]
        for sel in content_selectors:
            el = soup.select_one(sel) if ('.' in sel or '#' in sel or sel.startswith('#')) else soup.find(sel)
            if el:
                text = el.get_text(separator='\n', strip=True)
                text = re.sub(r'\s+', ' ', text).strip()
                if len(text) > 200 and not _is_garbage_text(text):
                    return text[:3000]
        body = soup.find('body')
        if body:
            text = body.get_text(separator=' ', strip=True)
            text = re.sub(r'\s+', ' ', text)
            if not _is_garbage_text(text):
                return text[:3000]
        return ""
    except Exception:
        return ""


# ===================== LLM 调用 =====================
def _call_llm(prompt: str, system_prompt: str = "", temperature: float = 0.3,
              max_tokens: int = 500) -> str:
    api_key = os.environ.get("DEEPSEEK_API_KEY", "")
    if not api_key:
        return ""
    try:
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": prompt})
        payload = {
            "model": LLM_MODEL,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
        }
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        resp = requests.post(LLM_API_URL, json=payload, headers=headers, timeout=30)
        if resp.status_code == 200:
            return resp.json()['choices'][0]['message']['content'].strip()
        return ""
    except Exception:
        return ""


def call_llm_summary(title: str, content: str, dept_name: str) -> str:
    if not content or len(content) < 50:
        return ""
    prompt = (
        f"你是一个专业的行业分析师，专注于{dept_name}领域。\n"
        f"请为以下新闻写一段100-200字的专业摘要，聚焦于行业影响和关键信息。\n\n"
        f"标题：{title}\n\n原文：\n{content[:2500]}\n\n"
        f"要求：摘要100-200字，突出行业影响和关键数据，用专业客观的语言。\n\n请写出摘要："
    )
    result = _call_llm(prompt, "你是一个专业的行业分析师。", max_tokens=500)
    if result:
        return result[:500]
    return content[:300].strip()


def call_llm_viewpoint_summary(articles_news: List[dict], articles_sentiment: List[dict],
                                dept_name: str) -> str:
    """生成核心观点摘要"""
    news_context = ""
    for i, a in enumerate(articles_news[:10], 1):
        news_context += f"{i}. [{a.get('info_type', '行业资讯')}] {a.get('title', '')}\n"
        s = a.get('summary', a.get('abstract', ''))
        if s: news_context += f"   摘要：{s[:100]}\n"
    sent_context = ""
    for a in articles_sentiment[:8]:
        cname = a.get('company', a.get('source_name', ''))
        sent_context += f"- {cname}: {a.get('title', '')}\n"
        s = a.get('summary', a.get('abstract', ''))
        if s: sent_context += f"  摘要：{s[:100]}\n"

    prompt = (
        f"你是一个专业的行业分析师，专注于{dept_name}领域。\n"
        f"请根据本周采集的以下资讯和舆情，撰写一份'核心观点摘要'，\n"
        f"从政策层面、市场层面、企业层面三个维度总结本周行业核心动态。\n\n"
        f"=== 行业资讯 ===\n{news_context}\n"
        f"=== 舆情信息 ===\n{sent_context}\n\n"
        f"要求：200-400字，分政策、市场、企业三个层面，每个层面写一段。\n"
        f"格式：每段以'1. 政策层面：'、'2. 市场层面：'、'3. 企业层面：'开头。\n\n"
        f"核心观点摘要："
    )
    result = _call_llm(prompt, "你是一个资深的行业分析师。", temperature=0.4, max_tokens=800)
    return result


def call_llm_conclusion_and_suggestions(articles_news: List[dict], articles_sentiment: List[dict],
                                         dept_name: str) -> Tuple[str, str]:
    """生成核心结论和工作建议"""
    news_context = ""
    for i, a in enumerate(articles_news[:8], 1):
        news_context += f"{i}. [{a.get('info_type', '行业资讯')}] {a.get('title', '')}\n"
        s = a.get('summary', a.get('abstract', ''))
        if s: news_context += f"   摘要：{s[:100]}\n"
    sent_context = ""
    for a in articles_sentiment[:8]:
        cname = a.get('company', a.get('source_name', ''))
        sent_context += f"- {cname}: {a.get('title', '')}\n"

    prompt = (
        f"你是一个资深的行业分析师，专注于{dept_name}领域。\n"
        f"请根据本周采集的资讯和舆情，撰写：\n"
        f"1. 核心结论（3-5条，总结本周行业整体判断、关键变量）\n"
        f"2. 工作建议（4-5条，按紧急/重要/常规/长期分级，含时间要求）\n\n"
        f"=== 行业资讯 ===\n{news_context}\n"
        f"=== 舆情信息 ===\n{sent_context}\n\n"
        f"格式要求：\n"
        f"结论部分以'核心结论：'开头，每条用'1.'、'2.'编号。\n"
        f"建议部分以'工作建议：'开头，每条格式为'【紧急/重要/常规/长期】工作事项 | 时间要求'\n\n"
        f"请输出："
    )
    text = _call_llm(prompt, "你是一个资深的行业分析师。", temperature=0.4, max_tokens=1000)
    if not text:
        return "", ""
    parts = text.split('工作建议：')
    conclusion = parts[0].replace('核心结论：', '').strip()
    suggestions = parts[1].strip() if len(parts) > 1 else ""
    return conclusion, suggestions


# ===================== 分类函数 =====================
def classify_info_type(title: str, abstract: str = "", source: str = "") -> str:
    combined = (title + ' ' + abstract + ' ' + source).lower()
    scores = {}
    for info_type, keywords in INFO_TYPE_KEYWORDS.items():
        scores[info_type] = sum(1 for kw in keywords if kw.lower() in combined)
    if not scores or max(scores.values()) == 0:
        return "行业资讯"
    return max(scores, key=scores.get)


def classify_sentiment_label(title: str, abstract: str = "") -> str:
    combined = (title + ' ' + abstract).lower()
    for label, keywords in SENTIMENT_LABEL_KEYWORDS.items():
        if any(kw.lower() in combined for kw in keywords):
            return label
    return "关注"


# ===================== DOCX 样式工具 =====================

def _set_run_font(run, size_pt=10, bold=False, color=COLOR_DARK, name=FONT_NAME):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def _add_empty_para(doc, before=0, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def _add_section_break(doc):
    """换页"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br)


def _add_section_title(doc, text):
    """章节标题：B45309 16pt 加粗 微软雅黑"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _set_run_font(run, 16, True, COLOR_ACCENT)
    return p


def _add_subsection_title(doc, text):
    """小节标题：1F2937 12pt 加粗 微软雅黑"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run_font(run, 12, True, COLOR_DARK)
    return p


def _set_cell_font(cell, size_pt=10, bold=False, color=COLOR_DARK, name=FONT_NAME, align=None):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            _set_run_font(run, size_pt, bold, color, name)


def _set_cell_padding(cell, top=60, bottom=60, left=100, right=100):
    """设置单元格内边距"""
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _set_cell_vertical_center(cell):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)


def _shade_cell(cell, hex_color):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _add_hyperlink(paragraph, url, text, size_pt=9, color="2563EB"):
    """在段落中添加可点击的超链接（兼容 WPS/Word）"""
    part = paragraph.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)

    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} {nsdecls("r")} r:id="{r_id}" w:history="1">'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:rStyle w:val="Hyperlink"/>'
        f'      <w:rFonts w:ascii="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
        f'      <w:color w:val="{color}"/>'
        f'      <w:sz w:val="{int(size_pt * 2)}"/>'
        f'      <w:szCs w:val="{int(size_pt * 2)}"/>'
        f'      <w:u w:val="none"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">{text}</w:t>'
        f'  </w:r>'
        f'</w:hyperlink>')
    paragraph._element.append(hyperlink)


def _make_table(doc, rows_data, headers, col_widths=None):
    """统一的表格生成函数"""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格边框
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # 表头行
    header_row = table.rows[0]
    for ci, header_text in enumerate(headers):
        cell = header_row.cells[ci]
        # 清空并设置文本
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header_text)
        _set_run_font(run, 10, True, COLOR_WHITE)
        _set_cell_padding(cell)
        _set_cell_vertical_center(cell)
        _shade_cell(cell, "1E3A5F")

    # 数据行
    last_col_is_link = any(h in ["来源链接"] for h in headers)
    link_col_idx = len(headers) - 1 if last_col_is_link else -1

    for ri, row_data in enumerate(rows_data):
        row = table.rows[ri + 1]
        for ci, cell_data in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # 如果是来源链接列，添加可点击的超链接
            if ci == link_col_idx and str(cell_data).startswith("http"):
                _add_hyperlink(p, str(cell_data), "点击查看")
            else:
                run = p.add_run(str(cell_data))
                _set_run_font(run, 10, False, COLOR_DARK)
            _set_cell_padding(cell)
            _set_cell_vertical_center(cell)

    # 设置列宽
    if col_widths:
        for ri in range(len(table.rows)):
            row = table.rows[ri]
            for ci in range(min(len(col_widths), len(row.cells))):
                row.cells[ci].width = Cm(col_widths[ci])

    return table


# ===================== 简报生成主函数 =====================

def _process_news_for_analysis(articles_news: List[dict], dept_name: str):
    """为行业影响分析生成政策/市场/企业三个维度的数据"""
    policy_items = []
    market_items = []
    enterprise_items = []
    for a in articles_news:
        info_type = classify_info_type(a.get('title', ''), a.get('summary', a.get('abstract', '')),
                                        a.get('source_name', ''))
        if info_type == "政策":
            policy_items.append(a)
        elif info_type == "市场":
            market_items.append(a)
        elif info_type in ("企业", "技术"):
            enterprise_items.append(a)
        else:
            enterprise_items.append(a)  # 默认归入企业
    return policy_items[:8], market_items[:8], enterprise_items[:8]


def _parse_suggestions_table(suggestions_text: str):
    """将LLM输出的建议文本解析为(优先级, 工作事项, 时间要求)列表"""
    rows = []
    if not suggestions_text:
        return rows
    for line in suggestions_text.strip().split('\n'):
        line = line.strip()
        if not line:
            continue
        # 匹配格式：【紧急】xxx | xxx 或 【重要】xxx|xxx
        m = re.match(r'[【\[]+(紧急|重要|常规|长期)[】\]]+(.*?)[|｜](.*)', line)
        if m:
            prio = m.group(1)
            task = m.group(2).strip()
            time_req = m.group(3).strip()
            rows.append((prio, task, time_req))
    if not rows:
        # 尝试松散匹配
        for line in suggestions_text.strip().split('\n'):
            line = line.strip()
            for prio in ["紧急", "重要", "常规", "长期"]:
                if prio in line:
                    parts = re.split(r'[|｜]', line)
                    task = parts[0].replace(f'【{prio}】', '').replace(f'[{prio}]', '').strip()
                    time_req = parts[1].strip() if len(parts) > 1 else "持续推进"
                    rows.append((prio, task if task else line, time_req))
                    break
    return rows


def _ensure_hyperlink_style(doc):
    """确保文档中包含 Hyperlink 样式，否则 WPS 可能不识别超链接"""
    styles = doc.styles
    for s in styles:
        if s.name == "Hyperlink":
            return
    # 如果不存在，创建一个
    from docx.enum.style import WD_STYLE_TYPE
    hyperlink_style = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    hyperlink_style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    hyperlink_style.font.underline = False
    hyperlink_style.font.name = FONT_NAME
    rPr = hyperlink_style._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), FONT_NAME)


def generate_brief_docx(dept_name: str, articles_news: List[dict],
                         articles_sentiment: List[dict], date_str: str,
                         skip_llm: bool = False) -> Document:
    """生成铝产业简报样式的DOCX"""
    doc = Document()

    # 确保 Hyperlink 样式存在
    _ensure_hyperlink_style(doc)

    # 页面设置（同参考文档）
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # 计算报告周期
    try:
        end_date = datetime.strptime(date_str, "%Y-%m-%d")
        start_date = end_date - timedelta(days=6)
        period = f"{start_date.strftime('%Y年%m月%d日')}-{end_date.strftime('%Y年%m月%d日')}"
    except:
        period = date_str

    # ==================== 封面 ====================
    _add_empty_para(doc, before=100, after=5)  # 上间距

    # 封面标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"{dept_name}行业")
    _set_run_font(run, 28, True, COLOR_ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("资讯分析简报")
    _set_run_font(run, 28, True, COLOR_DARK)

    _add_empty_para(doc, before=10, after=10)

    # 封面信息行
    info_items = [
        (f"报告周期：{period}", 11, COLOR_GRAY),
        ("报告类型：行业决策参考", 11, COLOR_GRAY),
        (f"发布日期：{date_str}", 11, COLOR_GRAY),
    ]
    for text, size, color in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(text)
        _set_run_font(run, size, False, color)

    # 数据来源
    sources = set()
    for a in articles_news:
        src = a.get('source_name', a.get('source', ''))
        if src and src not in ["百度搜索"]:
            sources.add(src)
    for a in articles_sentiment:
        src = a.get('source_name', a.get('source', ''))
        if src and src not in ["百度搜索"]:
            sources.add(src)
    source_str = "、".join(sorted(sources)[:10])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"数据来源：{source_str}")
    _set_run_font(run, 9, False, COLOR_GRAY)

    # ==================== 换页 ====================
    _add_section_break(doc)

    # ==================== 一、核心观点摘要 ====================
    _add_section_title(doc, "一、核心观点摘要")

    viewpoint_text = ""
    if not skip_llm:
        print(f"  🤖 生成核心观点摘要...")
        viewpoint_text = call_llm_viewpoint_summary(articles_news, articles_sentiment, dept_name)
        time.sleep(random.uniform(0.5, 1.0))

    if viewpoint_text:
        lines = viewpoint_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 按"X. XXX层面："拆分为标签+正文
            m = re.match(r'^(\d+\s*[.、]\s*[^：]+[：:])\s*(.*)', line)
            if m:
                label = m.group(1)
                body = m.group(2)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run_label = p.add_run(label)
                _set_run_font(run_label, 11, True, COLOR_ACCENT)
                run_body = p.add_run(body)
                _set_run_font(run_body, 10.5, False, COLOR_DARK)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(line)
                _set_run_font(run, 10.5, False, COLOR_DARK)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("(暂无核心观点摘要)")
        _set_run_font(run, 10.5, False, COLOR_GRAY)

    # ==================== 二、行业资讯 ====================
    news_count = len(articles_news)
    _add_section_title(doc, f"二、国内{dept_name}资讯汇总（{news_count}条）")

    news_rows = []
    for i, a in enumerate(articles_news, 1):
        info_type = classify_info_type(a.get('title', ''), a.get('summary', a.get('abstract', '')),
                                        a.get('source_name', ''))
        summary = a.get('summary', a.get('abstract', ''))[:120]
        title_full = a.get('title', '')
        core_content = f"{title_full}。{summary[:80]}" if summary else title_full
        source_name = a.get('source_name', a.get('source', ''))[:15]
        link = a.get('link', '')  # 不截断，完整URL
        news_rows.append((str(i), info_type, core_content, source_name, link or "#"))

    news_headers = ["序号", "咨询类型", "核心内容", "消息来源主体", "来源链接"]
    _make_table(doc, news_rows, news_headers, col_widths=[0.8, 1.2, 8.5, 2.5, 5.0])

    # ==================== 三、行业舆情（正向） ====================
    sentiment_positive = []
    sentiment_negative = []
    sentiment_other = []
    for a in articles_sentiment:
        label = classify_sentiment_label(a.get('title',''), a.get('summary',''))
        if label == "利好":
            sentiment_positive.append(a)
        elif label == "利空":
            sentiment_negative.append(a)
        else:
            sentiment_other.append(a)

    if sentiment_positive:
        _add_section_title(doc, f"三、全球行业正向舆情（{len(sentiment_positive)}条）")
        sent_rows = []
        for i, a in enumerate(sentiment_positive, 1):
            label = "利好"
            info_type = classify_info_type(a.get('title', ''), a.get('summary', a.get('abstract', '')),
                                            a.get('source_name', ''))
            summary = a.get('summary', a.get('abstract', ''))[:120]
            title_full = a.get('title', '')
            core_content = f"{title_full}。{summary[:80]}" if summary else title_full
            src_name = a.get('company', a.get('source_name', a.get('source', '')))[:15]
            link = a.get('link', '')  # 不截断
            sent_rows.append((str(i), label, info_type, core_content, src_name, link or "#"))
        t = _make_table(doc, sent_rows,
                        ["序号","舆情标签","咨询类型","核心内容","消息来源主体","来源链接"],
                        col_widths=[0.8, 1.0, 1.2, 8.0, 2.5, 4.0])
        # 给舆情标签列着色
        for ri in range(len(sent_rows)):
            label_cell = t.rows[ri+1].cells[1]
            _shade_cell(label_cell, COLOR_LIHOU)

    # ==================== 四、行业舆情（负面） ====================
    if sentiment_negative:
        _add_section_title(doc, f"四、全球行业负面舆情（{len(sentiment_negative)}条）")
        sent_rows = []
        for i, a in enumerate(sentiment_negative, 1):
            label = "利空"
            info_type = classify_info_type(a.get('title', ''), a.get('summary', a.get('abstract', '')),
                                            a.get('source_name', ''))
            summary = a.get('summary', a.get('abstract', ''))[:120]
            title_full = a.get('title', '')
            core_content = f"{title_full}。{summary[:80]}" if summary else title_full
            src_name = a.get('company', a.get('source_name', a.get('source', '')))[:15]
            link = a.get('link', '')  # 不截断
            sent_rows.append((str(i), label, info_type, core_content, src_name, link or "#"))
        t = _make_table(doc, sent_rows,
                        ["序号","舆情标签","咨询类型","核心内容","消息来源主体","来源链接"],
                        col_widths=[0.8, 1.0, 1.2, 8.0, 2.5, 4.0])
        for ri in range(len(sent_rows)):
            label_cell = t.rows[ri+1].cells[1]
            _shade_cell(label_cell, COLOR_LIKONG)

    # ==================== 五、本地企业舆情 ====================
    local_enterprise_items = [a for a in articles_sentiment if a.get('company','') and
                              classify_sentiment_label(a.get('title',''), a.get('summary','')) == "关注"]
    local_enterprise_items += [a for a in articles_sentiment if a.get('company','') and a not in local_enterprise_items]
    if not local_enterprise_items:
        local_enterprise_items = [a for a in articles_sentiment if a.get('company','')]

    if local_enterprise_items:
        # 按企业分组
        company_groups = {}
        for a in local_enterprise_items:
            cname = a.get('company', '其他')[:12]
            if cname not in company_groups:
                company_groups[cname] = []
            company_groups[cname].append(a)
        company_count = len(company_groups)
        _add_section_title(doc, f"五、贵阳本地企业舆情（{company_count}家企业）")

        ent_rows = []
        for a in local_enterprise_items:
            label = classify_sentiment_label(a.get('title',''), a.get('summary',''))
            info_type = classify_info_type(a.get('title',''), a.get('summary', a.get('abstract','')),
                                            a.get('source_name',''))
            cname = a.get('company', '其他')[:12]
            summary = a.get('summary', a.get('abstract', ''))[:120]
            title_full = a.get('title', '')
            clean_summary = _clean_summary(summary) if summary else ""
            core_content = f"{title_full}。{clean_summary[:90]}" if clean_summary else title_full
            src_name = a.get('source_name', a.get('source', ''))[:15]
            link = a.get('link', '')  # 不截断
            ent_rows.append((cname, label, info_type, core_content, src_name, link or "#"))
        t = _make_table(doc, ent_rows,
                        ["企业名称","舆情标签","咨询类型","核心内容","消息来源主体","来源链接"],
                        col_widths=[2.0, 1.0, 1.2, 7.5, 2.5, 4.0])
        for ri in range(len(ent_rows)):
            label_cell = t.rows[ri+1].cells[1]
            label_text = ent_rows[ri][1]
            if label_text == "利好":
                _shade_cell(label_cell, COLOR_LIHOU)
            elif label_text == "利空":
                _shade_cell(label_cell, COLOR_LIKONG)
            else:
                _shade_cell(label_cell, COLOR_GUANZHU)

    # ==================== 六、行业影响分析 ====================
    _add_section_title(doc, "六、行业影响分析")

    policy_items, market_items, enterprise_items = _process_news_for_analysis(articles_news, dept_name)

    # 6.1 政策影响分析
    _add_subsection_title(doc, "6.1 政策影响分析")
    pol_rows = []
    for a in policy_items[:6]:
        title_full = a.get('title', '')
        summary = a.get('summary', a.get('abstract', ''))[:100]
        impact = summary if summary else title_full[:80]
        next_impact = f"关注{title_full[:20]}后续进展"
        info_type_short = classify_info_type(title_full, summary, a.get('source_name',''))
        pol_rows.append((info_type_short, impact, next_impact))
    if not pol_rows:
        pol_rows = [("政策", "本周无重大政策资讯", "持续关注")]
    _make_table(doc, pol_rows,
                ["政策类型", "核心影响", "下一步行业影响"],
                col_widths=[3.0, 6.5, 6.5])

    # 6.2 市场影响分析
    _add_subsection_title(doc, "6.2 市场影响分析")
    mkt_rows = []
    for a in market_items[:6]:
        title_full = a.get('title', '')
        summary = a.get('summary', a.get('abstract', ''))[:100]
        impact = summary if summary else title_full[:80]
        next_impact = f"关注市场趋势"
        mkt_rows.append((title_full[:20], impact, next_impact))
    if not mkt_rows:
        mkt_rows = [("市场动态", "本周无重大市场资讯", "持续关注")]
    _make_table(doc, mkt_rows,
                ["影响因素", "核心影响", "下一步行业影响"],
                col_widths=[3.0, 6.5, 6.5])

    # 6.3 企业影响分析
    _add_subsection_title(doc, "6.3 企业影响分析")
    ent_rows2 = []
    if local_enterprise_items:
        # 按企业聚合
        company_agg = {}
        for a in local_enterprise_items:
            cname = a.get('company', '其他')[:12]
            if cname not in company_agg:
                company_agg[cname] = {"positive": [], "negative": [], "summary": ""}
            label = classify_sentiment_label(a.get('title',''), a.get('summary',''))
            title_short = a.get('title', '')[:20]
            if label == "利好":
                company_agg[cname]["positive"].append(title_short)
            else:
                company_agg[cname]["negative"].append(title_short)
        for cname, info in company_agg.items():
            pos = "、".join(info["positive"][:3]) if info["positive"] else "—"
            neg = "、".join(info["negative"][:3]) if info["negative"] else "—"
            suggestion = "持续关注"
            if info["positive"] and info["negative"]:
                suggestion = "两面兼顾，利用利好对冲风险"
            elif info["positive"]:
                suggestion = "把握机遇，扩大优势"
            elif info["negative"]:
                suggestion = "预警风险，提前应对"
            ent_rows2.append((cname, pos, neg, suggestion))
    if not ent_rows2:
        ent_rows2 = [("本地企业", "—", "—", "持续关注")]
    _make_table(doc, ent_rows2,
                ["企业名称", "正向影响", "负面/关注因素", "下一步建议"],
                col_widths=[2.5, 5.0, 5.0, 5.0])

    # ==================== 七、总结与建议 ====================
    _add_section_title(doc, "七、总结与建议")

    conclusion_text = ""
    suggestions_text = ""
    if not skip_llm:
        print(f"  🤖 生成总结与建议...")
        conclusion_text, suggestions_text = call_llm_conclusion_and_suggestions(
            articles_news, articles_sentiment, dept_name)
        time.sleep(random.uniform(0.5, 1.0))

    # 核心结论
    _add_subsection_title(doc, "核心结论")
    if conclusion_text:
        lines = conclusion_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            m = re.match(r'^(\d+\s*[.、])\s*(.*)', line)
            if m:
                num = m.group(1)
                body = m.group(2)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run_num = p.add_run(num + " ")
                _set_run_font(run_num, 11, True, COLOR_ACCENT)
                run_body = p.add_run(body)
                _set_run_font(run_body, 10.5, False, COLOR_DARK)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(line)
                _set_run_font(run, 10.5, False, COLOR_DARK)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("(暂无自动生成的结论)")
        _set_run_font(run, 10.5, False, COLOR_GRAY)

    # 工作建议
    _add_subsection_title(doc, "工作建议")

    suggestion_rows = _parse_suggestions_table(suggestions_text)
    if not suggestion_rows:
        # 如果LLM没解析出来，用默认
        suggestion_rows = [
            ("常规", "梳理本周资讯，形成阶段性总结", "每周例行工作"),
            ("常规", "根据政策动态调整关注重点", "根据实际需求"),
        ]

    t = _make_table(doc, suggestion_rows,
                    ["优先级", "工作事项", "时间要求"],
                    col_widths=[2.0, 8.5, 5.5])
    # 为优先级列着色
    for ri in range(len(suggestion_rows)):
        prio = suggestion_rows[ri][0]
        cell = t.rows[ri+1].cells[0]
        hex_color = PRIO_COLORS.get(prio, "F3F4F6")
        _shade_cell(cell, hex_color)

    # ==================== 底部 ====================
    _add_empty_para(doc, before=10, after=5)

    # 免责声明
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("免责声明：本报告基于公开信息整理，仅供行业决策参考，不构成投资建议。")
    _set_run_font(run, 9, False, COLOR_GRAY)

    # 编制单位
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("编制单位：智汇云数科技")
    _set_run_font(run, 9, False, COLOR_GRAY)

    # 数据来源
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"数据来源：{source_str}")
    _set_run_font(run, 9, False, COLOR_GRAY)

    return doc


# ===================== 处室处理 =====================

def process_dept(dept_name: str, all_articles: dict, date_str: str,
                 skip_llm: bool = False) -> Tuple[bool, str]:
    """处理单个处室的简报生成"""
    if dept_name not in all_articles:
        print(f"  ⚠️ {dept_name}: 没有数据，跳过")
        return False, ""

    articles = all_articles[dept_name]
    if not articles:
        print(f"  ⚠️ {dept_name}: 文章列表为空，跳过")
        return False, ""

    deduplicated = deduplicate_articles(articles)
    print(f"\n📝 {dept_name}: {len(articles)}条 → 去重后{len(deduplicated)}条")

    # 评分排序
    scored = [(score_article(a, dept_name), a) for a in deduplicated]
    scored.sort(key=lambda x: (-x[0], x[1].get('date', '')))

    # 区分新闻和舆情
    all_top = [a for _, a in scored[:20]]
    articles_news = []
    articles_sentiment = []
    for a in all_top:
        if a.get('company', ''):
            articles_sentiment.append(a)
        else:
            articles_news.append(a)

    # 如果新闻不足，补充
    if len(articles_news) < 10:
        for a in deduplicated:
            if a not in all_top and a['link'] not in {x['link'] for x in articles_news}:
                if not a.get('company', ''):
                    articles_news.append(a)
                if len(articles_news) >= 10:
                    break

    print(f"   精选: 新闻{len(articles_news)}条, 舆情{len(articles_sentiment)}条")

    if not skip_llm:
        # 获取原文 + 大模型摘要
        print(f"   🔍 获取原文并生成智能摘要...")
        summary_count = 0
        target_articles = articles_news + articles_sentiment
        for i, a in enumerate(target_articles):
            title_short = a.get('title', '')[:30]
            link = a.get('link', '')
            if link:
                print(f"     [{i+1}/{len(target_articles)}] {title_short}...", end=" ")
                body = fetch_article_text(link)
                if body:
                    summary = call_llm_summary(a['title'], body, dept_name)
                    if summary and not _is_garbage_text(summary) and len(summary) >= 30:
                        a['summary'] = _clean_summary(summary)
                        summary_count += 1
                        print(f"\u2705 ({len(summary)}字)")
                    else:
                        # LLM失败或返回垃圾 → 使用原始摘要（来自爬虫）
                        raw_abstract = a.get('abstract', a.get('summary', ''))[:300]
                        if raw_abstract and not _is_garbage_text(raw_abstract):
                            a['summary'] = _clean_summary(raw_abstract)
                            print(f"\u26a0\ufe0f (LLM失败，使用原始摘要)")
                        elif body and not _is_garbage_text(body):
                            a['summary'] = _clean_summary(body[:300])
                            print(f"\u26a0\ufe0f (LLM失败，截取正文)")
                        else:
                            a['summary'] = ""
                            print(f"\u26a0\ufe0f (LLM失败，原文不可用)")
                else:
                    raw_abstract = a.get('abstract', a.get('summary', ''))[:200]
                    if raw_abstract and not _is_garbage_text(raw_abstract):
                        a['summary'] = _clean_summary(raw_abstract)
                        print(f"\u26a0\ufe0f (无原文，使用原始摘要)")
                    else:
                        a['summary'] = ""
                        print(f"\u26a0\ufe0f (无原文且摘要不可用)")
                time.sleep(random.uniform(0.3, 0.8))
            else:
                a['summary'] = a.get('abstract', a.get('summary', ''))[:200]
                print(f"     [{i+1}/{len(target_articles)}] {title_short}: \u26a0\ufe0f (无链接)")
        print(f"    \u2705 生成 {summary_count}/{len(target_articles)} \u7bc7\u667a\u80fd\u6458\u8981")

    # 生成 .docx
    brief_dir = os.path.join(DATA_DIR, date_str, "brief")
    os.makedirs(brief_dir, exist_ok=True)
    docx_path = os.path.join(brief_dir, f"{dept_name}_模板简报_{date_str}.docx")

    doc = generate_brief_docx(dept_name, articles_news, articles_sentiment,
                               date_str, skip_llm)
    doc.save(docx_path)
    file_size = os.path.getsize(docx_path)
    print(f"  \u2705 \u6a21\u677f\u7b80\u62a5: {docx_path} ({file_size/1024:.1f} KB)")

    return True, docx_path


# ===================== 主入口 =====================

def main():
    import argparse
    parser = argparse.ArgumentParser(description="模板化简报生成 v3（铝产业简报样式）")
    parser.add_argument("--date", default=datetime.now().strftime("%Y-%m-%d"))
    parser.add_argument("--dept", default="all", help="指定处室（all为全部）")
    parser.add_argument("--skip-llm", action="store_true", help="跳过LLM摘要生成")
    args = parser.parse_args()

    date_str = args.date

    articles_path = os.path.join(DATA_DIR, date_str, "archive", "all_articles.json")
    if not os.path.exists(articles_path):
        print(f"\u26a0\ufe0f all_articles.json \u4e0d\u5b58\u5728: {articles_path}")
        print("请先运行采集脚本")
        sys.exit(1)

    with open(articles_path, encoding="utf-8") as f:
        all_articles = json.load(f)

    depts = list(all_articles.keys()) if args.dept == "all" else [args.dept]

    all_processed = []
    for dept_name in depts:
        success, docx_path = process_dept(dept_name, all_articles, date_str, args.skip_llm)
        if success:
            all_processed.append(docx_path)

    print(f"\n{'='*60}")
    print(f"\U0001f389 \u6a21\u677f\u7b80\u62a5\u751f\u6210\u5b8c\u6210!")
    for p in all_processed:
        print(f"  \U0001f4c4 {p}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()