from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime
import argparse
import os
import re


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = parse_xml(
        '<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'r:id="{}">'
        '<w:r>'
        '<w:rPr><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr>'
        '<w:t>{}</w:t>'
        '</w:r>'
        '</w:hyperlink>'.format(r_id, text)
    )
    paragraph._p.append(hyperlink)

# ── Reusable helpers from html_to_docx.py ──────────────────────────

def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
        fill
    ))
    tcPr.append(shading)

def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '</w:tcBorders>'
    )
    for edge, attrs in kwargs.items():
        element = parse_xml(
            '<w:{} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'w:val="{}" w:sz="{}" w:space="{}" w:color="{}"/>'.format(
                edge, attrs.get('val', 'single'), attrs.get('sz', '4'),
                attrs.get('space', '0'), attrs.get('color', '000000')
            )
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_paragraph_left_border(paragraph, color="165DFF", sz="12", space="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBorders = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:left w:val="single" w:sz="{}" w:space="{}" w:color="{}"/>'
        '</w:pBdr>'.format(sz, space, color)
    )
    pPr.append(pBorders)

def set_paragraph_top_border(paragraph, color="EEEEEE", sz="4", space="0"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBorders = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="{}" w:space="{}" w:color="{}"/>'
        '</w:pBdr>'.format(sz, space, color)
    )
    pPr.append(pBorders)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    twips = int(width_cm * 567)
    tcW = parse_xml(
        '<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:w="{}" w:type="dxa"/>'.format(twips)
    )
    tcPr.append(tcW)

def set_table_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        '<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblLayout = parse_xml(
        '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:type="fixed"/>'
    )
    tblPr.append(tblLayout)

def set_table_grid_cols(table, widths_cm):
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        return
    gridCols = tblGrid.findall(qn('w:gridCol'))
    for i, gridCol in enumerate(gridCols):
        if i < len(widths_cm):
            twips = int(widths_cm[i] * 567)
            gridCol.set(qn('w:w'), str(twips))
        else:
            gridCol.set(qn('w:w'), '1000')

def get_column_widths(headers):
    width_map = {
        '序号': 1.0,
        '舆情标签': 1.4,
        '企业名称': 2.0,
        '咨询类型': 1.8,
        '资讯类型': 1.8,
        '核心内容': 6.0,
        '消息来源主体': 1.8,
        '来源链接': 1.2,
        '政策类型': 2.5,
        '影响因素': 2.5,
        '核心影响': 4.5,
        '下一步行业影响': 4.5,
        '正向影响': 3.5,
        '负面/关注因素': 3.5,
        '下一步建议': 3.5,
        '优先级': 1.5,
        '工作事项': 7.0,
        '时间要求': 2.0,
    }
    widths = []
    for h in headers:
        matched = False
        for key, w in width_map.items():
            if key in h:
                widths.append(w)
                matched = True
                break
        if not matched:
            widths.append(2.5)
    return widths

# ── Markdown parsing ───────────────────────────────────────────────

def parse_md_table(lines):
    """Parse markdown table lines into list of row-cell lists."""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    # Drop separator row (second row usually contains :---)
    if len(rows) > 1:
        if all('---' in c or c == '' for c in rows[1]):
            rows = [rows[0]] + rows[2:]
    return rows

def is_instruction_line(line):
    """Skip Word operation hints / comments."""
    s = line.strip()
    if not s:
        return True
    if s.startswith('（') or s.startswith('('):
        return True
    if s.startswith('---'):
        return True
    if s == '此处换页':
        return True
    return False

def parse_md(path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    blocks = []
    in_cover = True
    table_buffer = []
    cover_info_lines = []

    for raw_line in lines:
        line = raw_line.rstrip('\n')
        if is_instruction_line(line):
            if line.strip() == '此处换页':
                # flush cover info before switching
                if cover_info_lines:
                    blocks.append({'type': 'cover_info', 'text': '\n'.join(cover_info_lines)})
                    cover_info_lines = []
                blocks.append({'type': 'page_break'})
                in_cover = False
            continue

        stripped = line.strip()

        # Table row
        if stripped.startswith('|'):
            table_buffer.append(stripped)
            continue
        else:
            if table_buffer:
                blocks.append({'type': 'table', 'rows': parse_md_table(table_buffer)})
                table_buffer = []

        if in_cover:
            if stripped.startswith('# ') and not stripped.startswith('## '):
                blocks.append({'type': 'cover_h1', 'text': stripped[2:]})
            elif stripped.startswith('## '):
                blocks.append({'type': 'cover_h2', 'text': stripped[3:]})
            else:
                cover_info_lines.append(stripped)
        else:
            if stripped.startswith('## '):
                blocks.append({'type': 'h3', 'text': stripped[3:]})
            elif stripped.startswith('### '):
                blocks.append({'type': 'h4', 'text': stripped[4:]})
            elif re.match(r'^\d+\.\s', stripped):
                blocks.append({'type': 'list_item', 'text': stripped})
            else:
                blocks.append({'type': 'paragraph', 'text': stripped})

    if table_buffer:
        blocks.append({'type': 'table', 'rows': parse_md_table(table_buffer)})
    if cover_info_lines:
        blocks.insert(0 if not blocks else next((i for i,b in enumerate(blocks) if b['type']!='cover_h1'), 0),
                      {'type': 'cover_info', 'text': '\n'.join(cover_info_lines)})

    return blocks

# ── Document builders ──────────────────────────────────────────────

def add_styled_paragraph(doc, text, base_size=10.5, bold_color='333333', is_summary=False):
    """Add paragraph with **bold** support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            color = 'E53935' if is_summary else bold_color
            run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
        else:
            run = p.add_run(part)
        run.font.size = Pt(base_size)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p

def add_table_block(doc, section, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    set_table_layout_fixed(table)

    headers = rows[0]
    col_widths = get_column_widths(headers)
    is_suggest = any('工作事项' in h for h in headers)

    # Scale to page width
    page_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    total_col_width = sum(col_widths)
    scale = page_width_cm / total_col_width if total_col_width > 0 else 1
    scaled_widths = [w * scale for w in col_widths]

    total_width_twips = int(page_width_cm * 567)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblW = parse_xml('<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{}" w:type="dxa"/>'.format(total_width_twips))
    tblPr.append(tblW)
    set_table_grid_cols(table, scaled_widths)

    # Column index maps for semantic coloring
    sentiment_col = next((i for i, h in enumerate(headers) if '舆情标签' in h), -1)
    priority_col = next((i for i, h in enumerate(headers) if '优先级' in h), -1)

    for i, row_cells in enumerate(rows):
        for j in range(col_count):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            is_header = (i == 0)
            cell_text = row_cells[j] if j < len(row_cells) else ''
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.0
            if not is_header and cell_text.startswith(('http://', 'https://')):
                add_hyperlink(p, '链接', cell_text)
            elif not is_header and cell_text.startswith(('[链接]')):
                res = re.search(r'\((http.*)\)', cell_text)
                if res:
                    add_hyperlink(p, '链接', res.group(1))
                else:
                    add_hyperlink(p, '链接', cell_text[5:-1])
            else:
                run = p.add_run(cell_text)
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                if is_header:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            if j < len(scaled_widths):
                set_cell_width(cell, scaled_widths[j])

            border_color = 'DDDDDD'
            set_cell_borders(cell,
                top={'val': 'single', 'sz': '4', 'color': border_color, 'space': '0'},
                bottom={'val': 'single', 'sz': '4', 'color': border_color, 'space': '0'},
                left={'val': 'single', 'sz': '4', 'color': border_color, 'space': '0'},
                right={'val': 'single', 'sz': '4', 'color': border_color, 'space': '0'},
            )

            if is_header:
                set_cell_shading(cell, '1E3A5F')
            else:
                val = cell_text.strip()
                # Sentiment labels
                if j == sentiment_col:
                    if val == '利好':
                        set_cell_shading(cell, '4CAF50')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif val == '利空':
                        set_cell_shading(cell, 'F44336')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif val == '关注':
                        set_cell_shading(cell, 'FF9800')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Priority labels
                if j == priority_col:
                    if val == '紧急':
                        set_cell_shading(cell, 'E53935')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif val == '重要':
                        set_cell_shading(cell, 'FB8C00')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif val == '常规':
                        set_cell_shading(cell, '1E88E5')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif val == '长期':
                        set_cell_shading(cell, '43A047')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

def convert_md_to_docx(md_path, docx_path):
    blocks = parse_md(md_path)
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    in_summary = False
    last_h3 = ''
    cover_subtitle = ''

    for blk in blocks:
        t = blk['type']

        if t == 'cover_h1':
            # Push cover content down by two blank lines
            for _ in range(4):
                doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(blk['text'])
            run.font.size = Pt(32)
            run.font.color.rgb = RGBColor(180, 83, 9)
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        elif t == 'cover_h2':
            cover_subtitle = blk['text']
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(20)
            run = p.add_run(blk['text'])
            run.font.size = Pt(28)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            # Divider line under h2
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(30)
            pPr = p2._p.get_or_add_pPr()
            pBorders = parse_xml(
                '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:bottom w:val="single" w:sz="1" w:space="1" w:color="B45309"/>'
                '</w:pBdr>'
            )
            pPr.append(pBorders)

        elif t == 'cover_info':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(blk['text'])
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        elif t == 'page_break':
            doc.add_page_break()

        elif t == 'h3':
            last_h3 = blk['text']
            in_summary = '核心观点摘要' in last_h3 or '核心结论' in last_h3
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(40)
            p.paragraph_format.space_after = Pt(20)
            run = p.add_run(blk['text'])
            run.font.size = Pt(16)
            run.font.color.rgb =  RGBColor(180, 83, 9)
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            #set_paragraph_left_border(p, color="165DFF", sz="12", space="8")

        elif t == 'h4':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(15)
            run = p.add_run(blk['text'])
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        elif t == 'list_item':
            add_styled_paragraph(doc, blk['text'], base_size=10.5, is_summary=in_summary)

        elif t == 'paragraph':
            # Disclaimer detection
            if '免责声明' in blk['text'] or '编制单位' in blk['text']:
                if '免责声明' in blk['text']:
                    # Two blank lines before disclaimer
                    for _ in range(2):
                        doc.add_paragraph()
                    # Separator line
                    p_sep = doc.add_paragraph()
                    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pPr_sep = p_sep._p.get_or_add_pPr()
                    pBorders_sep = parse_xml(
                        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:bottom w:val="single" w:sz="1" w:space="1" w:color="B45309"/>'
                        '</w:pBdr>'
                    )
                    pPr_sep.append(pBorders_sep)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(10)
                #set_paragraph_top_border(p, color='EEEEEE', sz='4', space='10')
                # Parse bold inside disclaimer too
                parts = re.split(r'(\*\*.*?\*\*)', blk['text'])
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            else:
                add_styled_paragraph(doc, blk['text'], base_size=10.5)

        elif t == 'table':
            add_table_block(doc, section, blk['rows'])

    # Header: cover subtitle right-aligned with bottom border
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(cover_subtitle)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Bottom border under header
    hpPr = hp._p.get_or_add_pPr()
    hpBorders = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="1" w:space="1" w:color="B45309"/>'
        '</w:pBdr>'
    )
    hpPr.append(hpBorders)

    # Footer: left = date + 行业决策参考, right = 第x页
    today = datetime.now().strftime('%Y年%m月%d日')
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Right-aligned tab stop at exact right margin
    usable_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(usable_width_cm), WD_TAB_ALIGNMENT.RIGHT)

    # Left text
    run = fp.add_run(f'{today} · 行业决策参考')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Tab to push page number to the right
    run_tab = fp.add_run('\t\t\t')
    run_tab.font.size = Pt(9)

    # Page number field: 第 PAGE 页
    run_pre = fp.add_run('第 ')
    run_pre.font.size = Pt(9)
    run_pre.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_pre.font.name = 'Microsoft YaHei'
    run_pre._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Insert PAGE field via XML
    run_pn = fp.add_run()
    fldChar_begin = parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>')
    run_pn._r.append(fldChar_begin)

    run_instr = fp.add_run()
    instrText = parse_xml('<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve"> PAGE </w:instrText>')
    run_instr._r.append(instrText)

    run_sep = fp.add_run()
    fldChar_sep = parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="separate"/>')
    run_sep._r.append(fldChar_sep)

    run_placeholder = fp.add_run('1')
    run_placeholder.font.size = Pt(9)
    run_placeholder.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_placeholder.font.name = 'Microsoft YaHei'
    run_placeholder._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    run_end = fp.add_run()
    fldChar_end = parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="end"/>')
    run_end._r.append(fldChar_end)

    run_post = fp.add_run(' 页')
    run_post.font.size = Pt(9)
    run_post.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_post.font.name = 'Microsoft YaHei'
    run_post._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX')
    parser.add_argument('--src', default='example.md', help='Input markdown filename (default: example.md)')
    parser.add_argument('--dst', default=None, help='Output docx filename (default: same basename as input with .docx)')
    parser.add_argument('--dstdir', default='.', help='Output directory path (default: current directory)')
    args = parser.parse_args()

    # Default output name matches input basename with .docx extension
    if args.dst is None:
        base, _ = os.path.splitext(args.src)
        args.dst = base + '.docx'

    out_path = os.path.join(args.dstdir, args.dst)
    convert_md_to_docx(args.src, out_path)
